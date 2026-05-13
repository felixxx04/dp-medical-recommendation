"""
毕设论文高级更新脚本 - 支持在指定位置插入段落
使用底层XML操作实现段落插入
"""
import os
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from copy import deepcopy

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
INPUT_FILE = os.path.join(SCRIPT_DIR, "thesis_updated_v3.docx")
OUTPUT_FILE = os.path.join(SCRIPT_DIR, "thesis_updated_v4.docx")

def set_run_font(run, font_name, size_pt, bold=False):
    """设置运行字体（含中文字体）"""
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

def create_paragraph_element(text, font_name="宋体", size_pt=12, bold=False,
                              alignment=None, first_line_indent=False, line_spacing=1.5):
    """创建段落XML元素"""
    p = OxmlElement('w:p')

    # 段落属性
    pPr = OxmlElement('w:pPr')

    # 对齐
    if alignment:
        jc = OxmlElement('w:jc')
        jc.set(qn('w:val'), alignment)
        pPr.append(jc)

    # 首行缩进
    if first_line_indent:
        ind = OxmlElement('w:ind')
        ind.set(qn('w:firstLine'), "420")  # 约0.74cm
        pPr.append(ind)

    # 行距
    if line_spacing:
        spacing = OxmlElement('w:spacing')
        spacing.set(qn('w:line'), str(int(line_spacing * 240)))
        pPr.append(spacing)

    p.append(pPr)

    # Run元素
    r = OxmlElement('w:r')

    # Run属性（字体）
    rPr = OxmlElement('w:rPr')
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:eastAsia'), font_name)
    rPr.append(rFonts)

    # 字号
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), str(int(size_pt * 2)))  # Word用半磅
    rPr.append(sz)
    szCs = OxmlElement('w:szCs')
    szCs.set(qn('w:val'), str(int(size_pt * 2)))
    rPr.append(szCs)

    # 粗体
    if bold:
        b = OxmlElement('w:b')
        rPr.append(b)

    r.append(rPr)

    # 文本
    t = OxmlElement('w:t')
    t.text = text
    r.append(t)

    p.append(r)
    return p

def insert_paragraph_after(doc, index, text, font_name="宋体", size_pt=12, bold=False,
                           alignment=None, first_line_indent=False, line_spacing=1.5):
    """在指定段落后插入新段落"""
    p_new = create_paragraph_element(text, font_name, size_pt, bold,
                                      alignment, first_line_indent, line_spacing)

    # 获取目标段落
    target_p = doc.paragraphs[index]._element
    target_p.addnext(p_new)

    return p_new

def add_chapter4_section0(doc):
    """在4.1之前插入4.0临床知识路由层"""
    print("正在插入 4.0 临床知识路由层...")

    # 查找4.1节位置
    idx_41 = -1
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if text.startswith("4.1") and ("安全" in text or "推荐" in text):
            idx_41 = i
            break

    if idx_41 == -1:
        print("  未找到4.1节位置")
        return False

    print(f"  找到4.1节位置: P{idx_41}")

    # 在4.1之前插入（即在idx_41-1之后插入）
    insert_idx = idx_41 - 1

    # 新增内容（从下往上插入，因为每次都在同一位置后插入）
    contents = [
        ("4.0.3 患者口语增强", "黑体", 12, False, None, False, None),  # 三级标题
        ("口语增强模块采用三层fallback策略处理患者非标准化输入：第一层精确匹配口语映射表中的完整短语；第二层关键词匹配提取输入中的医学关键词；第三层症状组合推理将多症状组合推断为最可能的疾病。当三层均失败时，降级为症状级别推荐并标记为低置信度。", "宋体", 12, False, None, True, 1.5),
        ("4.0.2 路由算法实现", "黑体", 12, False, None, False, None),
        ("路由层采用四级确定性路由：L1口语标准化将患者的非标准口语化中文表达标准化为标准英文医学术语，路由表包含154条口语到标准术语映射；L2疾病归类将标准化疾病名归类到身体系统和病因分类，路由表包含1495条疾病归类记录；L3适应症映射将身体系统与病因组合映射到对应的ATC药品治疗类别，路由表包含200条路由规则；L4药品召回在L3确定的药品类别范围内召回该类别中的所有候选药物。", "宋体", 12, False, None, True, 1.5),
        ("4.0.1 路由层设计理念", "黑体", 12, False, None, False, None),
        ("临床知识路由层是四层推荐架构的新增第零层，核心目标是解决中文疾病名到英文药品适应症之间的语义鸿沟。知识路由层通过四级确定性路由机制，在疾病名和药品之间建立了明确的多层路由网络，保证推荐的可解释性。", "宋体", 12, False, None, True, 1.5),
        ("4.0 临床知识路由层", "黑体", 14, True, None, False, None),  # 二级标题
    ]

    # 反向插入（从最后一行开始，这样插入后顺序正确）
    for text, font, size, bold, align, indent, spacing in reversed(contents):
        insert_paragraph_after(doc, insert_idx, text, font, size, bold, align, indent, spacing)

    print(f"  已插入4.0节（共{len(contents)}段）")
    return True

def add_chapter4_section5(doc):
    """在原4.5本章小结之前插入4.5审核反馈闭环"""
    print("正在插入 4.5 审核反馈闭环...")

    # 查找原4.5本章小结位置
    idx_summary = -1
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if "4.5" in text and "小结" in text:
            idx_summary = i
            break

    if idx_summary == -1:
        print("  未找到4.5本章小结位置")
        return False

    print(f"  找到4.5本章小结位置: P{idx_summary}")

    # 在小结之前插入（即在idx_summary-1之后）
    insert_idx = idx_summary - 1

    contents = [
        ("4.5.2 反馈学习机制", "黑体", 12, False, None, False, None),
        ("反馈学习器（FeedbackLearner）从review_log中读取审核决策，自动构建疾病到药品类别的惩罚权重图。当某药品类别被医生反复拒绝，系统自动对该配对施加渐进式惩罚，后续推荐中该药类候选药物得分乘以相应惩罚系数。当医生确认了曾被拒绝的药类时，惩罚系数逐步解除。", "宋体", 12, False, None, True, 1.5),
        ("4.5.1 审核流程设计", "黑体", 12, False, None, False, None),
        ("系统推荐2到4个候选药物后，医生进入审核环节，可进行三种操作：确认推荐（记录正面信号）、修改选择（从候选列表中另选药物）、拒绝推荐（标记为不适用）。审核决策通过ReviewPanel前端组件收集，经API接口写入MySQL review_log表。", "宋体", 12, False, None, True, 1.5),
        ("4.5 审核反馈闭环", "黑体", 14, True, None, False, None),  # 二级标题
    ]

    for text, font, size, bold, align, indent, spacing in reversed(contents):
        insert_paragraph_after(doc, insert_idx, text, font, size, bold, align, indent, spacing)

    print(f"  已插入4.5节（共{len(contents)}段）")

    # 更新原4.5小结为4.6
    print("  正在更新章节编号...")
    for i, para in enumerate(doc.paragraphs[idx_summary:idx_summary+5], start=idx_summary):
        if "4.5" in para.text:
            new_text = para.text.replace("4.5", "4.6")
            # 替换文本
            for run in para.runs:
                if "4.5" in run.text:
                    run.text = run.text.replace("4.5", "4.6")
            print(f"    P{i}: 4.5 → 4.6")
            break

    return True

def main():
    print("=" * 60)
    print("毕设论文高级更新脚本（支持段落插入）")
    print("=" * 60)
    print(f"输入文件: {INPUT_FILE}")
    print(f"输出文件: {OUTPUT_FILE}")
    print()

    if not os.path.exists(INPUT_FILE):
        print(f"错误: 输入文件不存在 {INPUT_FILE}")
        return

    print("正在加载文档...")
    doc = Document(INPUT_FILE)
    print(f"文档加载完成，共 {len(doc.paragraphs)} 个段落")
    print()

    # 插入新章节
    add_chapter4_section0(doc)
    add_chapter4_section5(doc)

    print()
    print("正在保存文档...")
    doc.save(OUTPUT_FILE)
    print(f"文档已保存至: {OUTPUT_FILE}")
    print()
    print("=" * 60)
    print("更新完成！")
    print("=" * 60)

if __name__ == "__main__":
    main()

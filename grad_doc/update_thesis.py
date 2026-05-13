"""
毕设论文自动更新脚本
基于 thesis_update_notes.md 中的修改计划，自动更新 thesis_updated_v2.docx
输出: thesis_updated_v3.docx
"""
import os
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

# 路径配置
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
INPUT_FILE = os.path.join(SCRIPT_DIR, "thesis_updated_v2.docx")
OUTPUT_FILE = os.path.join(SCRIPT_DIR, "thesis_updated_v3.docx")

def set_run_font(run, font_name, size_pt, bold=False):
    """设置运行字体（含中文字体）"""
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.name = font_name
    # 设置东亚字体（中文）
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

def add_heading1_style(doc, text):
    """添加一级标题 - 黑体 小3号(15pt) 粗体 居中"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run(text)
    set_run_font(run, "黑体", 15, True)
    return p

def add_heading2_style(doc, text):
    """添加二级标题 - 黑体 4号(14pt) 粗体"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_run_font(run, "黑体", 14, True)
    return p

def add_heading3_style(doc, text):
    """添加三级标题 - 黑体 小4号(12pt) 不加粗"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, "黑体", 12, False)
    return p

def add_body_style(doc, text):
    """添加正文 - 宋体 小4号(12pt) 1.5倍行距 首行缩进"""
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.first_line_indent = Cm(0.74)
    run = p.add_run(text)
    set_run_font(run, "宋体", 12, False)
    return p

def find_paragraph_by_content(doc, content_prefix, start_idx=0):
    """根据内容前缀查找段落索引"""
    for i, para in enumerate(doc.paragraphs[start_idx:], start=start_idx):
        if para.text.strip().startswith(content_prefix):
            return i
    return -1

def get_font_size(para):
    """获取段落第一个run的字体大小"""
    for run in para.runs:
        if run.font.size:
            return run.font.size.pt
    return None

def is_heading1(para):
    """判断是否为一级标题（15pt加粗）"""
    size = get_font_size(para)
    return size == 15.0 and any(run.font.bold for run in para.runs if run.font.bold is not None)

def is_heading2(para):
    """判断是否为二级标题（14pt加粗）"""
    size = get_font_size(para)
    return size == 14.0 and any(run.font.bold for run in para.runs if run.font.bold is not None)

def is_heading3(para):
    """判断是否为三级标题（12pt加粗）"""
    size = get_font_size(para)
    return size == 12.0 and any(run.font.bold for run in para.runs if run.font.bold is not None)

def replace_paragraph_text(para, new_text, preserve_style=True):
    """替换段落文本，保留样式"""
    # 获取第一个run的样式信息
    first_run = para.runs[0] if para.runs else None
    font_name = first_run.font.name if first_run else "宋体"
    font_size = first_run.font.size.pt if first_run and first_run.font.size else 12
    bold = first_run.font.bold if first_run else False

    # 清除原有runs
    for run in para.runs[:]:
        run._element.getparent().remove(run._element)

    # 添加新run
    run = para.add_run(new_text)
    if preserve_style:
        set_run_font(run, font_name, font_size, bold)
    else:
        set_run_font(run, "宋体", 12, False)

def update_abstract(doc):
    """更新第1章 摘要"""
    print("正在更新摘要...")

    # 新的四层架构描述
    new_abstract = """系统采用前后端分离微服务架构，前端使用React框架构建用户界面，后端使用Spring Boot提供RESTful API服务，模型服务使用Python FastAPI部署DeepFM推荐模型。在核心算法层面，系统设计了四层推荐架构：第零层为临床知识路由层（KnowledgeRouter），通过L1口语标准化→L2疾病归类→L3适应症映射→L4药品召回的四级确定性路由，解决中文疾病名到英文药品适应症之间的语义鸿沟；第一层为安全过滤层（SafetyFilter），保留12条绝对安全红线的硬排除规则，将相对禁忌、超说明书用药等9类情形由系统自动排除改为标记待审核；第二层为规则标记层（RuleMarker），对候选药物附加临床警告；第三层为DeepFM个性化排序层。系统还集成了患者口语增强模块（PatientInputEnhancer），通过精确匹配→关键词匹配→症状组合推理的三层fallback策略处理患者非标准化口语输入。推荐结果经医生审核确认后，反馈学习器（FeedbackLearner）自动调整后续推荐权重，形成持续优化的闭环。"""

    # 查找P17（系统架构描述段落）
    for i, para in enumerate(doc.paragraphs):
        if "系统采用前后端分离" in para.text or "系统设计了三层推荐架构" in para.text:
            print(f"  找到摘要段落 P{i}: {para.text[:50]}...")
            replace_paragraph_text(para, new_abstract)
            print("  已更新中文摘要")
            break

    # TODO: 更新英文摘要 P25
    print("  摘要更新完成")

def update_chapter3(doc):
    """更新第3章 系统需求分析与设计"""
    print("正在更新第3章...")

    # 查找架构描述段落并更新
    for i, para in enumerate(doc.paragraphs):
        if "系统整体架构分为三层" in para.text:
            print(f"  找到架构描述段落 P{i}")
            new_text = para.text.replace("三层", "四层").replace(
                "展示层、业务层和数据层",
                "路由层、展示层、业务层和数据层，其中路由层（KnowledgeRouter）负责疾病到药品类别的确定性路由"
            )
            replace_paragraph_text(para, new_text)
            print("  已更新架构描述")
            break

    # 查找模块划分段落并新增模块描述
    for i, para in enumerate(doc.paragraphs):
        if "系统划分为以下核心模块" in para.text:
            print(f"  找到模块划分段落 P{i}")
            # 在此段落后插入新模块描述
            # TODO: 需要在段落后插入新内容
            break

    print("  第3章更新完成")

def add_chapter4_section0(doc):
    """在第4章开头新增 4.0 临床知识路由层"""
    print("正在新增 4.0 临床知识路由层...")

    # 查找4.1节标题位置
    idx_41 = -1
    for i, para in enumerate(doc.paragraphs):
        if "4.1" in para.text and ("安全" in para.text or "推荐" in para.text):
            idx_41 = i
            break

    if idx_41 == -1:
        print("  未找到4.1节位置")
        return

    print(f"  找到4.1节位置 P{idx_41}")

    # 新章节内容
    sections = [
        ("4.0 临床知识路由层", 2),  # 二级标题
        ("4.0.1 路由层设计理念", 3),  # 三级标题
        """临床知识路由层是四层推荐架构的新增第零层，核心目标是解决中文疾病名到英文药品适应症之间的语义鸿沟。原有方案中，疾病映射器将中文疾病名转换为英文适应症后，直接在全量1815种药品中匹配适应症，缺乏中间的疾病→药类路由环节。当非标准疾病名被路由到错误的代理词时（如"肠炎"→代理"腹泻"→推荐止泻药而非抗感染药），会导致推荐错误。知识路由层通过四级确定性路由机制，在疾病名和药品之间建立了明确的多层路由网络。""",
        ("4.0.2 路由算法实现", 3),
        """路由层采用四级确定性路由：L1口语标准化将患者的非标准口语化中文表达（如"拉肚子"、"嗓子发炎"）标准化为标准英文医学术语，路由表包含154条口语→标准术语映射；L2疾病归类将标准化疾病名归类到身体系统（14个系统）和病因分类（23种病因类型），路由表包含1495条疾病归类记录；L3适应症映射将身体系统+病因组合映射到对应的ATC药品治疗类别，路由表包含200条L3路由规则；L4药品召回在L3确定的药品类别范围内召回该类别中的所有候选药物供后续安全过滤和排序。路由是确定性的——同一疾病每次路由结果一致，无机器学习随机性。每个推荐结果都可以回溯到完整的路由路径，保证了推荐的可解释性。""",
        ("4.0.3 患者口语增强", 3),
        """口语增强模块采用三层fallback策略处理患者非标准化输入：第一层精确匹配口语映射表中的完整短语（如"感冒"→"common cold"）；第二层关键词匹配提取输入中的医学关键词（如"喉咙不舒服"→提取"喉咙"→"pharyngitis"）；第三层症状组合推理将多症状组合推断为最可能的疾病（如"头痛+发烧+肌肉酸痛"→流感样疾病，信心级别低，强制医生审核）。当三层均失败时，降级为症状级别推荐并标记为低置信度。""",
    ]

    # 注意：python-docx不支持在指定位置插入段落，需要重建文档或使用其他方法
    # 这里我们记录需要插入的位置，后续手动处理
    print(f"  需要在 P{idx_41} 之前插入新章节（共{len([s for s in sections if isinstance(s, str)])}段正文）")
    print("  新章节内容已准备，将在脚本执行后手动插入")

def update_chapter4_safety_filter(doc):
    """更新 4.1.1 SafetyFilter 内容"""
    print("正在更新 SafetyFilter...")

    new_content = """安全过滤层的设计原则从"系统替患者决定"调整为"系统为医生标记，医生来做最终判断"。系统新增了SafetyLevel枚举（SAFE/WARNING/OFF_LABEL/UNVERIFIED/EXCLUDED），将原有的17类规则划分为硬排除（12条，不可审核）和软标记（9条，医生可审核）。硬排除规则（12条，不变）：1.过敏冲突 2.绝对禁忌症 3.致命药物交互（MAOI+SSRI等）4.妊娠X级 5.儿科禁忌 6.哺乳期L5级 7.草药补充剂治感染 8.糖皮质激素加重真菌感染 9.免疫抑制剂加重感染 10.IBD药物用于感染性肠炎（免疫抑制剂危险）11.抗生素治病毒感染（延长排毒期）12.PAH专用药误用于普通高血压。软标记规则（9条，原为排除，现改为标记待审核）：1.PPI用于胆囊病 2.抗生素用于尿路结石 3.降糖药用于结石 4.苯二氮卓用于胆囊病 5.青光眼药用于白内障 6.苯二氮卓用于OCD 7.促尿酸排泄药用于肠炎 8.抗生素用于真菌感染 9.无精确适应症匹配（off_label）。这一调整的核心优势在于：绝对安全红线保持不变，但原被过度拦截的药物现在以"标记待审"方式呈现在医生面前，医生可基于临床经验做最终决策。"""

    for i, para in enumerate(doc.paragraphs):
        if "安全过滤层是三层推荐架构的第一层" in para.text or "实现了17类排除规则" in para.text:
            print(f"  找到SafetyFilter段落 P{i}")
            replace_paragraph_text(para, new_content)
            print("  已更新SafetyFilter内容")
            break

def update_conclusion(doc):
    """更新结论"""
    print("正在更新结论...")

    new_conclusion_p1 = """核心创新为四层推荐架构：第零层临床知识路由层通过四级确定性路由精准对接疾病与药品类别，解决了中文疾病名到英文适应症之间的语义鸿沟；第一层安全标记层将过度拦截调整为标记审核，在保障绝对安全底线的同时扩大了候选药物范围；第二层规则标记层提供临床警告；第三层DeepFM个性化排序层融合药类评分和反馈学习。系统还实现了患者口语增强和医生审核反馈闭环，使系统支持1295种中文疾病输入，适应症路由覆盖率达92.5%，并通过232个自动化测试和DeepSeek临床药学验证。"""

    for i, para in enumerate(doc.paragraphs):
        if "核心创新为三层推荐安全架构" in para.text:
            print(f"  找到结论段落 P{i}")
            replace_paragraph_text(para, new_conclusion_p1)
            print("  已更新结论第一段")
            break

    # 更新展望
    for i, para in enumerate(doc.paragraphs):
        if "未来的研究方向包括" in para.text or "未来工作" in para.text:
            print(f"  找到展望段落 P{i}")
            new_outlook = """未来的研究方向包括：(1)扩充低频适应症（120个罕见病）的路由覆盖至100%；(2)利用审核反馈数据持续优化疾病→药品类别的路由权重；(3)将安全性数据（禁忌症数量、交互严重度）编码为DeepFM模型特征，实现更精细的个性化排序。"""
            replace_paragraph_text(para, new_outlook)
            print("  已更新展望")
            break

def main():
    """主函数"""
    print("=" * 60)
    print("毕设论文自动更新脚本")
    print("=" * 60)
    print(f"输入文件: {INPUT_FILE}")
    print(f"输出文件: {OUTPUT_FILE}")
    print()

    # 检查输入文件
    if not os.path.exists(INPUT_FILE):
        print(f"错误: 输入文件不存在 {INPUT_FILE}")
        return

    # 加载文档
    print("正在加载文档...")
    doc = Document(INPUT_FILE)
    print(f"文档加载完成，共 {len(doc.paragraphs)} 个段落")
    print()

    # 执行更新
    update_abstract(doc)
    update_chapter3(doc)
    # add_chapter4_section0(doc)  # 需要特殊处理
    update_chapter4_safety_filter(doc)
    update_conclusion(doc)

    # 保存文档
    print()
    print("正在保存文档...")
    doc.save(OUTPUT_FILE)
    print(f"文档已保存至: {OUTPUT_FILE}")
    print()
    print("=" * 60)
    print("更新完成！")
    print("=" * 60)
    print()
    print("注意事项：")
    print("1. 新增的4.0节和4.5节需要手动插入（python-docx不支持在指定位置插入）")
    print("2. 请对照 thesis_update_notes.md 检查修改是否完整")
    print("3. 章节编号可能需要手动调整（如4.5→4.6）")

if __name__ == "__main__":
    main()

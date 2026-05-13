"""
毕设论文高级更新脚本 v3 - 修复插入顺序
"""
import os
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
INPUT_FILE = os.path.join(SCRIPT_DIR, "thesis_updated_v2.docx")  # 从原始v2重新开始
OUTPUT_FILE = os.path.join(SCRIPT_DIR, "thesis_updated_final.docx")

def create_paragraph_element(text, font_name="宋体", size_pt=12, bold=False,
                              alignment=None, first_line_indent=False):
    """创建段落XML元素"""
    p = OxmlElement('w:p')
    pPr = OxmlElement('w:pPr')

    if alignment:
        jc = OxmlElement('w:jc')
        jc.set(qn('w:val'), alignment)
        pPr.append(jc)

    if first_line_indent:
        ind = OxmlElement('w:ind')
        ind.set(qn('w:firstLine'), "420")
        pPr.append(ind)

    # 行距1.5倍
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:line'), "360")
    pPr.append(spacing)

    p.append(pPr)

    r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:eastAsia'), font_name)
    rPr.append(rFonts)

    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), str(int(size_pt * 2)))
    rPr.append(sz)
    szCs = OxmlElement('w:szCs')
    szCs.set(qn('w:val'), str(int(size_pt * 2)))
    rPr.append(szCs)

    if bold:
        b = OxmlElement('w:b')
        rPr.append(b)

    r.append(rPr)

    t = OxmlElement('w:t')
    t.text = text
    r.append(t)

    p.append(r)
    return p

def insert_paragraphs_before(doc, index, paragraphs_data):
    """在指定位置之前插入多个段落（顺序插入）"""
    target_p = doc.paragraphs[index]._element

    # 按顺序插入（第一个先插入，然后依次在其后插入）
    for i, (text, font, size, bold, align, indent) in enumerate(paragraphs_data):
        p_new = create_paragraph_element(text, font, size, bold, align, indent)
        if i == 0:
            target_p.addprevious(p_new)
            last_inserted = p_new
        else:
            last_inserted.addnext(p_new)
            last_inserted = p_new

def update_abstract(doc):
    """更新摘要"""
    print("更新摘要...")

    new_abstract = """系统采用前后端分离微服务架构，前端使用React框架构建用户界面，后端使用Spring Boot提供RESTful API服务，模型服务使用Python FastAPI部署DeepFM推荐模型。在核心算法层面，系统设计了四层推荐架构：第零层为临床知识路由层（KnowledgeRouter），通过L1口语标准化→L2疾病归类→L3适应症映射→L4药品召回的四级确定性路由，解决中文疾病名到英文药品适应症之间的语义鸿沟；第一层为安全过滤层（SafetyFilter），保留12条绝对安全红线的硬排除规则，将相对禁忌、超说明书用药等9类情形由系统自动排除改为标记待审核；第二层为规则标记层（RuleMarker），对候选药物附加临床警告；第三层为DeepFM个性化排序层。系统还集成了患者口语增强模块（PatientInputEnhancer），通过精确匹配→关键词匹配→症状组合推理的三层fallback策略处理患者非标准化口语输入。推荐结果经医生审核确认后，反馈学习器（FeedbackLearner）自动调整后续推荐权重，形成持续优化的闭环。"""

    for i, para in enumerate(doc.paragraphs):
        if "系统采用前后端分离" in para.text or "系统设计了三层推荐架构" in para.text:
            for run in para.runs:
                run.text = ""
            run = para.add_run(new_abstract)
            run.font.size = Pt(12)
            run.font.name = "宋体"
            run._element.rPr.rFonts.set(qn('w:eastAsia'), "宋体")
            print(f"  已更新摘要 P{i}")
            break

def update_architecture(doc):
    """更新架构描述"""
    print("更新架构描述...")

    for i, para in enumerate(doc.paragraphs):
        if "系统整体架构分为三层" in para.text:
            new_text = para.text.replace("三层", "四层").replace(
                "展示层、业务层和数据层",
                "路由层、展示层、业务层和数据层，其中路由层（KnowledgeRouter）负责疾病到药品类别的确定性路由"
            )
            for run in para.runs:
                run.text = ""
            run = para.add_run(new_text)
            run.font.size = Pt(12)
            run.font.name = "宋体"
            run._element.rPr.rFonts.set(qn('w:eastAsia'), "宋体")
            print(f"  已更新架构描述 P{i}")
            break

def update_safety_filter(doc):
    """更新SafetyFilter"""
    print("更新SafetyFilter...")

    new_content = """安全过滤层的设计原则从"系统替患者决定"调整为"系统为医生标记，医生来做最终判断"。系统将原有的17类规则划分为硬排除（12条，不可审核）和软标记（9条，医生可审核）。硬排除规则包括：过敏冲突、绝对禁忌症、致命药物交互、妊娠X级、儿科禁忌、哺乳期L5级等12条绝对安全红线。软标记规则包括：PPI用于胆囊病、抗生素用于尿路结石等9类可审核情形。这一调整的核心优势在于：绝对安全红线保持不变，但原被过度拦截的药物现在以"标记待审"方式呈现，医生可基于临床经验做最终决策。"""

    for i, para in enumerate(doc.paragraphs):
        if "安全过滤层是三层推荐架构的第一层" in para.text or "实现了17类排除规则" in para.text:
            for run in para.runs:
                run.text = ""
            run = para.add_run(new_content)
            run.font.size = Pt(12)
            run.font.name = "宋体"
            run._element.rPr.rFonts.set(qn('w:eastAsia'), "宋体")
            print(f"  已更新SafetyFilter P{i}")
            break

def add_section_4_0(doc):
    """添加4.0临床知识路由层"""
    print("添加4.0临床知识路由层...")

    # 查找4.1位置
    idx_41 = -1
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip().startswith("4.1") and "安全" in para.text:
            idx_41 = i
            break

    if idx_41 == -1:
        print("  未找到4.1位置")
        return

    print(f"  在P{idx_41}之前插入")

    # 按正确顺序定义内容
    contents = [
        ("4.0 临床知识路由层", "黑体", 14, True, "center", False),
        ("4.0.1 路由层设计理念", "黑体", 12, False, None, False),
        ("临床知识路由层是四层推荐架构的新增第零层，核心目标是解决中文疾病名到英文药品适应症之间的语义鸿沟。知识路由层通过四级确定性路由机制，在疾病名和药品之间建立了明确的多层路由网络，保证推荐的可解释性。", "宋体", 12, False, None, True),
        ("4.0.2 路由算法实现", "黑体", 12, False, None, False),
        ("路由层采用四级确定性路由：L1口语标准化将患者的非标准口语化中文表达标准化为标准英文医学术语，路由表包含154条映射；L2疾病归类将标准化疾病名归类到身体系统和病因分类，路由表包含1495条记录；L3适应症映射将身体系统与病因组合映射到对应的ATC药品治疗类别，路由表包含200条规则；L4药品召回在L3确定的药品类别范围内召回该类别中的所有候选药物。", "宋体", 12, False, None, True),
        ("4.0.3 患者口语增强", "黑体", 12, False, None, False),
        ("口语增强模块采用三层fallback策略处理患者非标准化输入：第一层精确匹配口语映射表中的完整短语；第二层关键词匹配提取输入中的医学关键词；第三层症状组合推理将多症状组合推断为最可能的疾病。当三层均失败时，降级为症状级别推荐并标记为低置信度。", "宋体", 12, False, None, True),
    ]

    insert_paragraphs_before(doc, idx_41, contents)
    print(f"  已插入{len(contents)}段")

def add_section_4_5(doc):
    """添加4.5审核反馈闭环并更新原4.5为4.6"""
    print("添加4.5审核反馈闭环...")

    # 查找原4.5本章小结
    idx_summary = -1
    for i, para in enumerate(doc.paragraphs):
        if "4.5" in para.text and "小结" in para.text:
            idx_summary = i
            break

    if idx_summary == -1:
        print("  未找到4.5本章小结")
        return

    print(f"  在P{idx_summary}之前插入")

    contents = [
        ("4.5 审核反馈闭环", "黑体", 14, True, "center", False),
        ("4.5.1 审核流程设计", "黑体", 12, False, None, False),
        ("系统推荐2到4个候选药物后，医生进入审核环节，可进行三种操作：确认推荐（记录正面信号）、修改选择（从候选列表中另选药物）、拒绝推荐（标记为不适用）。审核决策通过ReviewPanel前端组件收集，经API接口写入MySQL review_log表。", "宋体", 12, False, None, True),
        ("4.5.2 反馈学习机制", "黑体", 12, False, None, False),
        ("反馈学习器（FeedbackLearner）从review_log中读取审核决策，自动构建疾病到药品类别的惩罚权重图。当某药品类别被医生反复拒绝，系统自动对该配对施加渐进式惩罚（1次拒绝乘0.7，2次乘0.5，3次及以上乘0.3）。当医生确认了曾被拒绝的药类时，惩罚系数逐步解除。", "宋体", 12, False, None, True),
    ]

    insert_paragraphs_before(doc, idx_summary, contents)
    print(f"  已插入{len(contents)}段")

    # 更新章节编号
    for para in doc.paragraphs:
        if "4.5本章小结" in para.text or ("4.5" in para.text and "小结" in para.text):
            for run in para.runs:
                run.text = run.text.replace("4.5", "4.6")
            print("  已更新章节编号 4.5→4.6")
            break

def update_conclusion(doc):
    """更新结论"""
    print("更新结论...")

    new_conclusion = """核心创新为四层推荐架构：第零层临床知识路由层通过四级确定性路由精准对接疾病与药品类别，解决了中文疾病名到英文适应症之间的语义鸿沟；第一层安全标记层将过度拦截调整为标记审核，在保障绝对安全底线的同时扩大了候选药物范围；第二层规则标记层提供临床警告；第三层DeepFM个性化排序层融合药类评分和反馈学习。系统还实现了患者口语增强和医生审核反馈闭环，使系统支持1295种中文疾病输入，适应症路由覆盖率达92.5%，并通过232个自动化测试和DeepSeek临床药学验证。"""

    for i, para in enumerate(doc.paragraphs):
        if "核心创新为三层推荐安全架构" in para.text:
            for run in para.runs:
                run.text = ""
            run = para.add_run(new_conclusion)
            run.font.size = Pt(12)
            run.font.name = "宋体"
            run._element.rPr.rFonts.set(qn('w:eastAsia'), "宋体")
            print(f"  已更新结论 P{i}")
            break

    # 更新展望
    new_outlook = """未来的研究方向包括：(1)扩充低频适应症（120个罕见病）的路由覆盖至100%；(2)利用审核反馈数据持续优化疾病→药品类别的路由权重；(3)将安全性数据（禁忌症数量、交互严重度）编码为DeepFM模型特征，实现更精细的个性化排序。"""

    for i, para in enumerate(doc.paragraphs):
        if "未来的研究方向包括" in para.text or "未来工作" in para.text:
            for run in para.runs:
                run.text = ""
            run = para.add_run(new_outlook)
            run.font.size = Pt(12)
            run.font.name = "宋体"
            run._element.rPr.rFonts.set(qn('w:eastAsia'), "宋体")
            print(f"  已更新展望 P{i}")
            break

def main():
    print("=" * 60)
    print("毕设论文更新脚本 v3")
    print("=" * 60)
    print(f"输入: {INPUT_FILE}")
    print(f"输出: {OUTPUT_FILE}")
    print()

    doc = Document(INPUT_FILE)
    print(f"加载完成，共 {len(doc.paragraphs)} 段落\n")

    # 执行所有更新
    update_abstract(doc)
    update_architecture(doc)
    update_safety_filter(doc)
    add_section_4_0(doc)
    add_section_4_5(doc)
    update_conclusion(doc)

    doc.save(OUTPUT_FILE)
    print(f"\n保存完成: {OUTPUT_FILE}")
    print(f"最终段落总数: {len(doc.paragraphs)}")

if __name__ == "__main__":
    main()

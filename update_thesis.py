#!/usr/bin/env python3
"""Update thesis document to reflect model service refactoring."""
import sys
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.shared import Pt
from copy import deepcopy

doc = Document('grad_doc/thesis_with_screenshots.docx')

def replace_para_text(para, new_text):
    """Replace paragraph text while preserving its style and run formatting."""
    style = para.style
    # Get font info from first run if exists
    font_name = None
    font_size = None
    if para.runs:
        r0 = para.runs[0]
        font_name = r0.font.name
        font_size = r0.font.size
    # Clear all runs
    for run in para.runs:
        run.text = ""
    # Set text via first run or add new run
    if para.runs:
        para.runs[0].text = new_text
    else:
        run = para.add_run(new_text)
    # Restore font info
    if para.runs and font_name:
        para.runs[0].font.name = font_name
    if para.runs and font_size:
        para.runs[0].font.size = font_size
    # Preserve style
    para.style = style

# ============================================================
# ABSTRACT UPDATE (Para 16-18)
# ============================================================
replace_para_text(doc.paragraphs[17],
    '系统采用前后端分离的微服务架构，前端基于React框架构建交互界面，后端使用Spring Boot提供RESTful API服务，模型服务则通过FastAPI部署深度学习推理接口。在核心算法层面，系统设计了三层推荐架构：第一层为安全过滤层（SafetyFilter），通过确定性规则排除绝对禁忌药物，确保推荐的临床安全性，差分隐私噪声不影响此层；第二层为规则标记层（RuleMarker），对相对禁忌和中等风险药物添加警示标记而非直接排除；第三层为个性化排序层，基于DeepFM深度因子分解机模型对安全候选药物进行精准评分，差分隐私噪声仅在此层应用。系统还集成了临床匹配器（ClinicalMatcher）实现标准化的疾病-适应症匹配，疾病映射器（DiseaseMapper）支持中文症状到英文疾病编码的语义转换，以及解释生成器（ExplanationGenerator）为推荐结果提供可解释性支撑。')

replace_para_text(doc.paragraphs[18],
    '实验结果表明，系统在设置合理隐私预算（ε=0.1）的条件下，推荐准确率保持在可接受范围内（下降约10%），同时有效降低了敏感信息泄露风险。差分隐私机制对系统性能的影响较小（约5毫秒），适合在实际医疗场景中部署应用。三层推荐架构确保了绝对禁忌药物不会被噪声意外推荐，临床阈值后处理（score<0.15置零、ceiling=min(1.0, raw+0.35)）防止噪声过度扭曲推荐方向。本系统为隐私保护型医疗推荐应用提供了兼顾临床安全性与个性化精准度的技术参考与实践借鉴。')

# ============================================================
# CHAPTER 3 UPDATES
# ============================================================

# 3.2.1 整体架构 (Para 143) - Add three-layer architecture description
replace_para_text(doc.paragraphs[143],
    '系统整体架构分为三层：展示层、业务层和数据层。展示层负责用户界面的呈现和交互，包括Web前端应用，基于React框架开发。业务层负责业务逻辑的处理，包括后端API服务和模型推理服务，后端服务基于Spring Boot框架开发，模型服务基于FastAPI框架开发。数据层负责数据的持久化存储，包括MySQL关系数据库和向量数据库。在业务层内部，模型推理服务采用了三层推荐架构作为核心设计：安全过滤层（SafetyFilter）通过确定性规则排除绝对禁忌药物；规则标记层（RuleMarker）对相对禁忌药物添加警示标记；个性化排序层基于DeepFM模型对安全候选药物进行精准评分并应用差分隐私噪声。三层架构确保了临床安全性与隐私保护的分离——差分隐私噪声仅作用于排序层，不会影响安全过滤的确定性判断。')

# 3.2.2 模块划分 (Para 145) - Add new modules
replace_para_text(doc.paragraphs[145],
    '根据功能需求，系统划分为以下核心模块：（1）认证模块：处理用户注册、登录、权限验证等操作，使用JWT实现无状态认证。（2）患者管理模块：管理患者的基本信息和健康档案，提供增删改查接口，支持V2生理指标字段（肾功能、肝功能、BMI、妊娠状态、哺乳状态、吸烟饮酒状况、血压、空腹血糖、糖化血红蛋白、胆固醇、心率等）。（3）药物管理模块：管理药物数据库，包括药物基本信息、适应症、禁忌症、副作用、相互作用等专业数据，以JSON格式存储多值字段。（4）安全过滤模块（SafetyFilter）：实现三层推荐架构的第一层，通过9类确定性排除规则（绝对禁忌症、儿科禁忌、过敏冲突、重大药物相互作用、妊娠X类、哺乳L5级、MAOI+SSRI禁忌、肾功能/肝功能严重损害排除、草药补充剂排除）过滤不安全药物，差分隐私噪声不影响此层。（5）规则标记模块（RuleMarker）：实现三层推荐架构的第二层，对相对禁忌、中等相互作用、妊娠C/D类警告、肾功能/肝功能轻度警告等添加requires_review和safetyType标记。（6）临床匹配模块（ClinicalMatcher）：实现标准化的疾病-适应症匹配，支持indication_match_conditions优先匹配、过敏标准化匹配和疾病语义扩展匹配。（7）疾病映射模块（DiseaseMapper）：支持中文症状/疾病名称到英文疾病编码的语义转换，包括SEMANTIC_VOCAB_MAP处理词汇表未覆盖的疾病映射。（8）推荐排序模块：基于DeepFM模型对安全候选药物进行个性化评分，应用差分隐私噪声和临床阈值后处理。（9）解释生成模块（ExplanationGenerator）：为推荐结果生成适应症匹配详情、安全性分析和推荐理由说明。（10）隐私管理模块：管理差分隐私配置、隐私预算追踪（基于强组合定理）和隐私事件日志。（11）审计日志模块：记录完整的推荐审计轨迹，支持知情同意日志记录。')

# 3.3.2 患者表设计 (Para 151) - Add v2 physiological fields
replace_para_text(doc.paragraphs[151],
    '患者表（patient）存储患者的基本信息，包括患者ID（主键）、姓名、性别、出生日期、联系电话、创建时间等字段。健康档案表（health_record）存储患者的疾病史、过敏史、当前用药等健康信息，以JSON格式存储便于扩展。在V2版本中，健康档案表新增了12项生理指标字段：肾功能等级（renal_function）、肝功能等级（hepatic_function）、BMI值、妊娠状态（pregnancy_status）、哺乳状态（breastfeeding_status）、吸烟状态（smoking_status）、饮酒状态（drinking_status）、血压（blood_pressure）、空腹血糖（fasting_glucose）、糖化血红蛋白（hba1c）、胆固醇（cholesterol）、心率（heart_rate）。这些生理指标字段为三层推荐架构中的安全过滤和规则标记提供了关键的临床决策依据，肾功能和肝功能等级直接影响SafetyFilter的排除规则和RuleMarker的警告标记。')

# 3.4.2 模型服务接口 (Para 161) - Update with new endpoints
replace_para_text(doc.paragraphs[161],
    '模型服务提供以下API接口：预测接口POST /model/predict接收患者特征（包括V2生理指标），返回三层推荐架构处理后的推荐药物列表，包含安全过滤排除药物、规则标记警告药物和个性化排序推荐药物。模型状态接口GET /model/status返回模型加载状态、设备信息、药物数量等。药物加载接口POST /model/load-drugs接收后端药物数据，加载到模型服务并构建安全数据映射（禁忌症映射95.3%覆盖、相互作用映射81.5%覆盖、总体安全数据97.0%覆盖）。隐私预算接口包括GET /model/privacy/budget获取预算状态和POST /model/privacy/budget/reset重置预算。审计日志接口包括GET /model/audit/logs查询审计记录和POST /model/audit/consent记录知情同意。训练接口POST /model/train支持Focal Loss训练参数（focalLossAlpha=0.25、focalLossGamma=2.0）。')

# ============================================================
# CHAPTER 4 - CORE ALGORITHMS (Major rewrite)
# ============================================================

# 4 intro (Para 165)
replace_para_text(doc.paragraphs[165],
    '本章详细阐述医疗用药推荐系统的核心算法设计与实现，包括三层推荐安全架构、DeepFM推荐模型、差分隐私机制以及推荐流程集成等核心部分。三层推荐架构是系统设计的核心创新，将推荐过程划分为安全过滤层（SafetyFilter）、规则标记层（RuleMarker）和个性化排序层，确保推荐结果在隐私保护的前提下兼顾临床安全性和个性化精准度。安全过滤层通过确定性规则排除绝对禁忌药物，差分隐私噪声不影响此层；规则标记层对相对禁忌药物添加警示标记而非直接排除；个性化排序层基于DeepFM模型对安全候选药物进行精准评分，差分隐私噪声仅在此层应用。这种分层设计既保障了临床安全性不受隐私噪声干扰，又实现了隐私保护与个性化推荐的有机融合。')

# 4.1 title: DeepFM推荐模型 → 三层推荐安全架构
replace_para_text(doc.paragraphs[166], '4.1  三层推荐安全架构')

# 4.1 intro (Para 167)
replace_para_text(doc.paragraphs[167],
    '三层推荐安全架构是本系统核心算法设计的基石，将推荐过程严格划分为三个层次，确保临床安全性与隐私保护的分离。该架构的设计理念源于医疗推荐场景的特殊性：在医疗用药推荐中，安全性是首要约束，绝对禁忌药物的推荐可能导致严重临床后果，因此安全性判断必须基于确定性规则，不能受差分隐私噪声影响。三层架构的核心原则是：差分隐私噪声仅作用于个性化排序层，安全过滤层和规则标记层均基于确定性临床规则，确保推荐结果的临床安全性不受隐私保护机制的干扰。')

# 4.1.1 title: 模型架构设计 → 安全过滤层
replace_para_text(doc.paragraphs[168], '4.1.1  安全过滤层（SafetyFilter）')

# 4.1.1 content (Para 169)
replace_para_text(doc.paragraphs[169],
    '安全过滤层是三层推荐架构的第一层，负责通过确定性规则排除绝对禁忌药物，确保推荐候选集中不含任何可能对患者造成严重危害的药物。安全过滤层实现了9类排除规则：（1）绝对禁忌症排除：药物禁忌症与患者疾病匹配时直接排除，使用ClinicalMatcher标准化匹配indication_match_conditions优先；（2）儿科禁忌排除：患者年龄<18时排除标记为pediatric_contraindicated的药物；（3）过敏冲突排除：患者过敏史与药物过敏警告匹配时排除；（4）重大药物相互作用排除：患者当前用药与候选药物存在重大相互作用时排除；（5）妊娠X类排除：妊娠患者排除Pregnancy Category X药物；（6）哺乳L5级排除：哺乳患者排除Lactation Safety L5药物；（7）MAOI+SSRI禁忌排除：MAOI类药物与SSRI类药物组合排除；（8）肾功能严重损害排除：严重肾功能不全患者排除肾功能禁忌药物；（9）肝功能严重损害排除：严重肝功能不全患者排除肝功能禁忌药物。安全过滤层的判断结果不受差分隐私噪声影响，确保绝对禁忌药物不会被噪声意外推入推荐候选集。')

# 4.1.2 title: FM组件实现 → 规则标记层
replace_para_text(doc.paragraphs[170], '4.1.2  规则标记层（RuleMarker）')

# 4.1.2 content (Para 171)
replace_para_text(doc.paragraphs[171],
    '规则标记层是三层推荐架构的第二层，对安全过滤层未排除的候选药物进行软标记，添加requires_review和safetyType标记而非直接排除。规则标记层实现了以下标记规则：（1）相对禁忌症标记：药物存在相对禁忌症时标记requires_review=True、safetyType=relative_contraindication；（2）中等药物相互作用标记：存在中等程度相互作用时标记requires_review=True、safetyType=moderate_interaction；（3）妊娠C/D类警告标记：妊娠患者使用C/D类药物时标记requires_review=True、safetyType=pregnancy_warning；（4）肾功能警告标记：肾功能轻度异常患者使用特定药物时标记肾功能相关警告；（5）肝功能警告标记：肝功能轻度异常患者使用特定药物时标记肝功能相关警告；（6）生育力警告标记：可能影响生育力的药物添加fertility_warning标记；（7）数据未验证标记：安全数据缺失的药物标记data_unverified。规则标记层的关键设计原则是：标记而非排除。标记后的药物仍保留在候选集中参与个性化排序，但携带警示信息供医生参考。这种设计避免了对相对禁忌药物的过度排除，在保障安全性的同时保留了个性化推荐的灵活性。')

# 4.1.3 title: 深度组件实现 → 临床匹配器
replace_para_text(doc.paragraphs[172], '4.1.3  临床匹配器（ClinicalMatcher）')

# 4.1.3 content (Para 173)
replace_para_text(doc.paragraphs[173],
    '临床匹配器是三层推荐架构的关键支撑组件，实现了标准化的疾病-适应症匹配算法，替代了早期的子串匹配方式，显著提升了匹配准确性。ClinicalMatcher提供三种核心匹配功能：（1）match_condition：标准化的疾病匹配，支持indication_match_conditions优先匹配。药物适应症数据包含两种字段：indication_match_conditions（精确匹配条件列表）和indications（原始文本列表）。匹配时优先使用indication_match_conditions进行精确匹配，仅在缺失时回退到indications的子串匹配。同时支持疾病语义扩展匹配，通过SEMANTIC_VOCAB_MAP将"甲减"等非标准疾病名称映射到"hypothyroidism"等标准英文编码。（2）match_allergy：标准化的过敏匹配，将患者过敏信息与药物过敏警告进行精确比对。（3）match_interaction：药物相互作用匹配，检查患者当前用药与候选药物的相互作用等级。临床匹配器的引入解决了早期子串匹配存在的误匹配问题（如"低血压"子串匹配到"高血压"相关药物），将匹配准确率从约70%提升至95%以上。')

# 4.2 title: 差分隐私机制实现 → DeepFM推荐模型
replace_para_text(doc.paragraphs[174], '4.2  DeepFM推荐模型')

# 4.2 intro (Para 175)
replace_para_text(doc.paragraphs[175],
    'DeepFM模型是三层推荐架构中个性化排序层的核心算法，负责对安全过滤层过滤后、规则标记层标记后的候选药物进行精准评分。本系统采用DeepFM v3版本，相比原始DeepFM模型进行了多项架构优化：合并嵌入设计、连续特征旁路、LayerNorm替代BatchNorm、差异化Dropout以及Focal Loss训练支持。')

# 4.2.1 title: 推理阶段差分隐私 → 模型架构设计
replace_para_text(doc.paragraphs[176], '4.2.1  模型架构设计')

# 4.2.1 content (Para 177)
replace_para_text(doc.paragraphs[177],
    'DeepFM v3模型的输入特征包括14个类别字段（age_group、gender、bmi_group、renal_function、hepatic_function、primary_disease、secondary_disease、allergy_severity、drug_class、med_class_1至med_class_4、pregnancy_cat、rx_otc、drug_candidate）和4个连续特征（age_raw、bmi_raw、gfr_raw、liver_score_raw）。类别特征首先经过合并嵌入层（Merged Embedding）转换为低维稠密向量，该层使用单个nn.Embedding实例配合field_offsets寄存器缓冲区实现Opacus兼容的差分隐私训练。嵌入维度为8，总嵌入空间大小为各字段词汇量之和。合并嵌入的优势在于：统一的参数管理、差分隐私训练兼容性（Opacus要求单一nn.Embedding）、以及内存效率优化。连续特征通过旁路机制直接绕过嵌入层和FM/Deep组件，以原始数值形式与最终输出拼接，避免了连续特征经过嵌入层的信息损失。这种设计使得模型能够同时学习类别特征的交互关系和连续特征的直接效应。')

# 4.2.2 title: 训练阶段差分隐私 → FM组件与深度组件
replace_para_text(doc.paragraphs[178], '4.2.2  FM组件与深度组件实现')

# 4.2.2 content (Para 179)
replace_para_text(doc.paragraphs[179],
    'FM组件（MultiFieldFM）基于合并嵌入实现，使用field_offsets寄存器缓冲区将不同字段的嵌入向量从统一嵌入表中提取。FM组件计算一阶线性效应和二阶特征交互，二阶交互通过嵌入向量的内积计算特征两两之间的交互强度。FM组件引入了embed_dropout机制，在训练时对嵌入向量进行随机丢弃以增强泛化能力。深度组件（Deep）采用多层全连接神经网络结构，隐藏层维度为[64, 32]，每层后接LayerNorm和ReLU激活函数。相比原始DeepFM使用BatchNorm，LayerNorm在推理时无需依赖批统计量，更适合单样本推理的医疗推荐场景。深度组件还实现了差异化Dropout策略：第一层Dropout率为0.3（较高，防止过拟合），第二层Dropout率为0.1（较低，保留更多有用信息）。DeepFM v3模型输出为原始logits值，sigmoid激活函数在推理时手动应用，这种设计支持Focal Loss训练中的概率计算需求。')

# 4.2.3 title: 隐私预算管理 → 连续特征旁路与Focal Loss
replace_para_text(doc.paragraphs[180], '4.2.3  连续特征旁路与Focal Loss训练')

# 4.2.3 content (Para 181)
replace_para_text(doc.paragraphs[181],
    '连续特征旁路是DeepFM v3的关键架构创新。4个连续特征（age_raw、bmi_raw、gfr_raw、liver_score_raw）不经过嵌入层处理，而是通过独立的线性变换层直接映射到输出空间。旁路机制的数学表达为：continuous_output = W_cont × continuous_input + b_cont，其中W_cont为权重矩阵，b_cont为偏置向量。连续特征旁路与FM输出、Deep输出在最终层拼接，形成完整的模型输出：output = FM_output + Deep_output + continuous_output。这种设计使得连续特征（特别是肾功能GFR和肝功能评分等关键生理指标）的信息不被嵌入层压缩损失，直接参与最终评分计算。在训练方面，系统支持Focal Loss训练策略，通过调整focalLossAlpha（默认0.25）和focalLossGamma（默认2.0）参数，使模型在训练时更加关注难以分类的样本。Focal Loss的数学定义为：FL(p) = -alpha × (1-p)^gamma × log(p)，其中alpha平衡正负样本权重，gamma降低易分类样本的损失权重。Focal Loss在医疗推荐场景中特别有用，因为正样本（适合药物）比例远低于负样本（不适合药物），需要模型更关注少数正样本的学习。')

# 4.3 title: RAG检索增强 → 差分隐私机制实现
replace_para_text(doc.paragraphs[182], '4.3  差分隐私机制实现')

# 4.3 intro (Para 183)
replace_para_text(doc.paragraphs[183],
    '为保护患者隐私，系统在个性化排序层实现了差分隐私保护机制。差分隐私噪声仅作用于三层推荐架构的第三层（个性化排序层），不影响安全过滤层和规则标记层的确定性判断。推理阶段噪声注入包括临床阈值后处理和隐私预算管理两个核心机制。')

# 4.3.1 title: 向量数据库构建 → 推理阶段差分隐私
replace_para_text(doc.paragraphs[184], '4.3.1  推理阶段差分隐私')

# 4.3.1 content (Para 185)
replace_para_text(doc.paragraphs[185],
    '在推理阶段，系统向DeepFM模型的推荐得分添加噪声实现隐私保护。支持的噪声机制包括拉普拉斯噪声和高斯噪声。拉普拉斯噪声尺度参数为Δf/ε（Δf为敏感度），高斯噪声方差σ = Δf × √(2ln(1.25/δ)) / ε。差分隐私噪声应用后，系统执行临床阈值后处理以确保推荐结果的临床合理性：（1）临床安全阈值：得分低于0.15的药物直接置零。此阈值是公开的确定性参数，根据差分隐私后处理定理，确定性后处理操作不降低隐私保护强度。0.15阈值确保了低分药物不会被噪声意外提升到推荐位置，同时保留了对安全过滤层排除药物的双重保障。（2）天花板截断：推荐得分上限为min(1.0, raw_score + 0.35)，防止噪声将低分药物放大超过3.5倍，避免噪声过度扭曲推荐方向。（3）异常检测：标记dpAnomaly=True当噪声显著改变推荐排序方向时，帮助医生识别受噪声影响较大的推荐结果。（4）置信区间计算：为每个推荐得分计算95%置信区间（拉普拉斯噪声CI=±2b/√3，高斯噪声CI=±1.96σ），当推荐药物的置信区间重叠时标记uncertainRanking=True，提示排序不确定性。')

# 4.3.2 title: 检索流程 → 隐私预算管理
replace_para_text(doc.paragraphs[186], '4.3.2  隐私预算管理')

# 4.3.2 content (Para 187)
replace_para_text(doc.paragraphs[187],
    '系统提供隐私预算管理功能，基于强组合定理（Strong Composition Theorem）追踪隐私预算消耗。强组合定理相比基础组合定理提供了更精确的隐私预算计算：对于k次ε-差分隐私操作的组合，总隐私预算为ε_total = ε√(2kln(1/δ值)) + kε(e^ε-1) + δ，其中δ值为目标松弛参数。这种计算方式使得在相同隐私保护强度下允许更多次推荐操作，提升了系统的实用性。系统实现了BudgetWarningLevel三级预警机制：normal（预算充足）、warning（预算消耗超过50%）、exceeded（预算耗尽）。当隐私预算耗尽时，系统拒绝新的推荐请求并提示用户重置预算，防止隐私保护强度不足。隐私配置参数包括：隐私预算ε（范围0.01~10.0）、松弛参数δ（范围0<δ<1）、敏感度Δf（范围0.01~1.0）、噪声机制类型（laplace/gaussian/geometric）。默认配置为ε=0.1的高斯噪声机制，适用于医疗数据的高敏感性特点。系统同时实现了双重隐私预算追踪：后端MySQL持久化存储（privacy_config.budget_used + privacy_ledger）和模型服务内存追踪（PrivacyBudgetTracker），两者独立运行以满足不同场景需求。')

# 4.3.3 title: 推荐理由生成 → 推荐质量保障机制
replace_para_text(doc.paragraphs[188], '4.3.3  推荐质量保障机制')

# 4.3.3 content (Para 189)
replace_para_text(doc.paragraphs[189],
    '系统实现了多层推荐质量保障机制，确保推荐结果的可靠性和临床实用性。质量保障机制包括以下核心组件：（1）最低可靠得分阈值（MIN_RELIABLE_SCORE=0.3）：推荐得分低于0.3的药物标记为低可靠性，提示医生谨慎参考。此阈值高于临床安全阈值0.15，形成了双重保障：0.15以下直接排除，0.15-0.3标记为低可靠性。（2）最低区分度检查（MIN_SEPARATION=0.15）：当推荐药物之间的得分差异小于0.15时，标记排序不确定性，提示推荐区分度不足。（3）置信区间重叠检测：当推荐药物的95%置信区间存在重叠时标记uncertainRanking=True，提示差分隐私噪声可能影响了排序结果的可靠性。（4）质量警告等级：系统定义了NO_RELIABLE_RECOMMENDATION（无可靠推荐）和LOW_CONFIDENCE（低置信度）两种质量警告状态，当所有推荐得分均低于0.3或置信区间大面积重叠时触发。（5）跨候选药物相互作用检查（DDI Cross-Candidate Check）：在最终推荐列表中检查推荐药物之间的相互作用，若发现重大相互作用则添加interactionWarning标记。该检查通过critical_interactions.py中定义的DDI配对数据库实现，覆盖了MAOI+SSRI、华法林+NSAIDs等高风险药物组合。')

# 4.4 intro (Para 191)
replace_para_text(doc.paragraphs[191],
    '系统将三层推荐安全架构、DeepFM模型、差分隐私机制和解释生成模块进行集成，形成完整的推荐流程。推荐流程严格遵循三层架构的层次化执行顺序，确保临床安全性、隐私保护和个性化精准度的有机融合。')

# 4.4.1 title: 推荐流程设计 → 三层推荐流程设计
replace_para_text(doc.paragraphs[192], '4.4.1  三层推荐流程设计')

# 4.4.1 content (Para 193)
replace_para_text(doc.paragraphs[193],
    '三层推荐流程按照以下步骤严格执行：第一步获取患者信息，包括基本信息、疾病史、过敏史、当前用药和V2生理指标（肾功能、肝功能、BMI等）；第二步通过DiseaseMapper将中文疾病名称映射为英文疾病编码，使用SEMANTIC_VOCAB_MAP处理词汇表未覆盖的疾病，区分primary_input_diseases（患者原始输入疾病）和vocab_diseases（词汇表可识别的代理疾病）；第三步执行安全过滤层（SafetyFilter），通过9类确定性排除规则过滤绝对禁忌药物，缩小候选集范围；第四步执行规则标记层（RuleMarker），对安全候选药物添加相对禁忌和中等风险标记；第五步基于DeepFM模型对安全候选药物进行个性化评分，应用差分隐私噪声和临床阈值后处理（0.15阈值置零、0.35天花板截断）；第六步通过疾病平衡选择算法（_select_disease_balanced）确保推荐药物覆盖多种疾病，优先覆盖lost_diseases（词汇表未覆盖的真实疾病）；第七步通过ExplanationGenerator为推荐结果生成适应症匹配详情、安全性分析和推荐理由；第八步执行跨候选DDI检查，检测推荐药物之间的相互作用风险；第九步应用质量保障机制检查推荐可靠性和置信度，返回最终推荐结果。')

# 4.4.2 title: 特征预处理 → 疾病映射与特征预处理
replace_para_text(doc.paragraphs[194], '4.4.2  疾病映射与特征预处理')

# 4.4.2 content (Para 195)
replace_para_text(doc.paragraphs[195],
    '疾病映射模块（DiseaseMapper）负责将患者输入的中文疾病名称转换为模型可识别的英文疾病编码。疾病映射面临的核心挑战是中文疾病表述的多样性和非标准化：患者可能使用"甲减"、"甲状腺功能减退"、"thyroid issues"等不同表述指代同一疾病。DiseaseMapper通过SEMANTIC_VOCAB_MAP解决这一问题，该映射表将常见的中文疾病简称映射到标准英文编码，如"甲减"→hypothyroidism、"高血压"→hypertension、"糖肾"→diabetic_nephropathy等。对于SEMANTIC_VOCAB_MAP仍无法覆盖的疾病，系统标记为lost_diseases，在推荐选择时优先覆盖这些疾病对应的适应症药物。特征预处理模块将原始患者信息和药物信息转换为DeepFM模型的输入格式。14个类别特征通过Pipeline Schema定义的字段映射编码为离散值，4个连续特征（age_raw、bmi_raw、gfr_raw、liver_score_raw）以原始数值形式通过旁路机制输入模型。药物候选字段（drug_candidate）覆盖1815种药物，每种药物作为独立候选进行评分。')

# 4.4.3 title: 结果后处理 → 推荐解释生成与结果后处理
replace_para_text(doc.paragraphs[196], '4.4.3  推荐解释生成与结果后处理')

# 4.4.3 content (Para 197)
replace_para_text(doc.paragraphs[197],
    '推荐解释生成由ExplanationGenerator模块实现，为每个推荐药物生成三层解释信息：（1）适应症匹配详情：说明药物适应症与患者疾病的匹配情况，使用ClinicalMatcher的indication_match_conditions进行精确匹配描述，区分直接匹配和语义扩展匹配；（2）安全性分析：综合SafetyFilter排除原因和RuleMarker标记信息，说明药物的安全性评估结果，包括是否存在相对禁忌、中等相互作用或妊娠/肾功能/肝功能警告；（3）推荐理由说明：综合适应症匹配和安全性分析，生成可理解的推荐理由文本，说明为何推荐该药物以及需要注意的事项。推荐结果经过后处理后返回给用户，后处理包括：药物名称英中翻译（通过DrugTranslator和TranslationMapper实现，包括drugName中文名和englishName英文名）、跨候选DDI相互作用检查、质量保障等级判定、置信区间展示和差分隐私噪声量记录。系统还实现了药物类别、安全类型、副作用等专业术语的全字段中文翻译，确保前端展示的一致性和可理解性。')

# 4.5 本章小结 (Para 199)
replace_para_text(doc.paragraphs[199],
    '本章详细阐述了医疗用药推荐系统的核心算法设计与实现。首先设计了三层推荐安全架构，包括安全过滤层（SafetyFilter）的9类确定性排除规则、规则标记层（RuleMarker）的软标记机制和临床匹配器（ClinicalMatcher）的标准化匹配算法，确保了临床安全性不受差分隐私噪声影响。然后设计了DeepFM v3推荐模型架构，实现了合并嵌入层、连续特征旁路、LayerNorm+差异化Dropout和Focal Loss训练支持，提升了模型的推理效率和训练效果。接着实现了推理阶段差分隐私机制，包括临床阈值后处理（0.15安全阈值、0.35天花板截断）、异常检测和95%置信区间计算，以及基于强组合定理的隐私预算管理和多层质量保障机制。最后集成了三层推荐流程、疾病映射、特征预处理、推荐解释生成和跨候选DDI检查，形成了完整的推荐处理链路。本章内容为系统实现提供了坚实的算法基础，确保了推荐结果在隐私保护前提下的临床安全性和个性化精准度。')

# ============================================================
# CHAPTER 5.4 UPDATE (Para 243)
# ============================================================
replace_para_text(doc.paragraphs[243],
    '模型服务基于FastAPI框架开发（v2.0.0版本），负责三层推荐架构的执行、DeepFM模型推理和差分隐私保护。项目主要包括models、services、data、utils等模块。models模块包含DeepFM v3模型定义（MultiFieldFM合并嵌入、Deep深度组件、DeepFM集成模型），使用PyTorch框架实现，支持Opacus差分隐私训练兼容。services模块包含RecommendationPredictor（三层推荐架构执行、疾病平衡选择、质量保障）、SafetyFilter（9类确定性排除规则）、RuleMarker（软标记规则）、ExplanationGenerator（推荐解释生成）四个核心服务。data模块包含critical_interactions.py（重大DDI配对数据库）、pipeline_data.json（禁忌症映射95.3%覆盖、相互作用映射81.5%覆盖、合并药物数据）。utils模块包含clinical_matcher.py（标准化疾病-适应症匹配）、disease_mapper.py（中文疾病语义映射）、drug_translator.py（药物名称英中翻译）、translation_mapper.py（统一翻译映射）、privacy.py（拉普拉斯/高斯噪声实现）、privacy_budget.py（强组合定理预算追踪）、audit_logger.py（审计日志系统）。模型服务实现了结构化异常层级（ModelServiceError、DataNotFoundError、DataValidationError、ModelNotLoadedError、PredictionError、PrivacyBudgetExceededError等），提供清晰的错误分类和HTTP状态码映射。模型输入采用V2版PredictRequest，支持14个类别特征和12项生理指标字段。')

# ============================================================
# CONCLUSION UPDATE (Para 259)
# ============================================================
replace_para_text(doc.paragraphs[259],
    '本文针对医疗用药推荐过程中的隐私保护问题，设计并实现了一种基于差分隐私保护的AI驱动个性化医疗用药推荐系统。系统采用前后端分离的微服务架构，核心创新为三层推荐安全架构，将推荐过程划分为安全过滤层（SafetyFilter）、规则标记层（RuleMarker）和个性化排序层，确保了临床安全性不受差分隐私噪声影响。在个性化排序层融合了DeepFM推荐模型和差分隐私机制，同时集成了临床匹配器、疾病映射器和解释生成器，在保障患者隐私的前提下实现了精准的个性化用药推荐。')

# Para 260 - main work and innovations
replace_para_text(doc.paragraphs[260],
    '本系统的主要工作和创新点包括：第一，设计了三层推荐安全架构，将推荐过程严格划分为安全过滤层、规则标记层和个性化排序层，确保差分隐私噪声仅作用于排序层而不影响安全过滤的确定性判断，解决了隐私保护与临床安全性之间的冲突。第二，将差分隐私机制与深度学习推荐模型相结合，在推理阶段实现了临床阈值后处理（0.15安全阈值、0.35天花板截断）和95%置信区间计算，基于后处理定理保证了隐私保护强度不被后处理操作降低。第三，设计了临床匹配器（ClinicalMatcher）实现标准化的疾病-适应症匹配，替代子串匹配方式，将匹配准确率提升至95%以上。第四，设计了疾病映射器（DiseaseMapper）支持中文疾病语义转换，通过SEMANTIC_VOCAB_MAP处理词汇表未覆盖的疾病映射。第五，实现了推荐解释生成器，为推荐结果提供适应症匹配详情、安全性分析和推荐理由说明，提升了推荐的可解释性。')

# Para 261 - testing results
replace_para_text(doc.paragraphs[261],
    '系统测试结果表明，在设置合理隐私预算（ε=0.1）的条件下，系统能够有效保护患者隐私，同时保持较高的推荐准确率。三层推荐架构确保了绝对禁忌药物不会被差分隐私噪声意外推荐，临床阈值后处理防止了噪声对推荐方向的过度扭曲。差分隐私机制对系统性能的影响较小（约5毫秒），适合在实际医疗场景中部署应用。')

# ============================================================
# SAVE
# ============================================================
doc.save('grad_doc/thesis_updated.docx')
print("Thesis updated successfully! Saved to grad_doc/thesis_updated.docx")
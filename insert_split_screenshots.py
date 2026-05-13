"""
Insert split screenshots into thesis docx.
Each image at 18.25cm width, ≤20cm doc height.
Figure numbering: single pages get 图5-X, split pages get 图5-Xa/5-Xb/5-Xc.
"""
import sys
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os

doc_path = 'grad_doc/初稿(1).docx'
screenshots_dir = 'grad_doc/screenshots'
output_path = 'grad_doc/thesis_with_screenshots.docx'

pages = [
    {
        'name': '系统首页',
        'desc': '系统首页是用户进入系统后的第一个页面，展示了系统的核心定位——差分隐私保护的智能用药推荐。页面顶部导航栏提供首页、患者档案、隐私配置、用药推荐、效果可视化、后台管理等功能入口。首页主体区域包含系统介绍标语"精准用药推荐，守护患者隐私"，以及"开始用药推荐"、"效果可视化"、"管理后台"三个快捷入口按钮。下方展示系统核心数据指标（隐私保护等级ε≤1.0、推荐准确率92%+、药物种类5,000+、服务患者10,000+）和四大核心功能模块（差分隐私保护、深度学习推荐、个性化用药、安全可控）的简介卡片。界面如图5-1所示。',
        'images': [('final_homepage_1.png', '图5-1  系统首页界面')]
    },
    {
        'name': '登录页面',
        'desc': '登录页面提供用户身份认证入口。页面左侧展示系统Logo和名称"智医荐药——基于差分隐私保护的智能用药推荐平台"，右侧为登录表单，包含账号和密码输入框以及"登录系统"按钮。页面底部提供测试账号信息，方便开发测试。系统采用JWT令牌进行身份认证，支持管理员（admin）、医生（doctor）和研究员（researcher）三种角色登录，不同角色拥有不同的功能权限。界面如图5-2所示。',
        'images': [('final_login_1.png', '图5-2  登录页面界面')]
    },
    {
        'name': '用药推荐页面',
        'desc': '用药推荐页面是系统的核心功能页面，为选定患者生成个性化用药推荐。页面顶部显示推荐功能标题和简介。用户首先通过患者选择器选择目标患者，系统自动展示患者的基本信息（姓名、年龄、性别）和健康档案（慢性疾病、过敏信息、当前用药）。页面提供双输入模式，用户可在疾病/症状输入框中输入自定义症状描述，与患者档案中的疾病信息结合进行推荐。隐私参数配置区域允许用户设置隐私预算ε、松弛参数δ、噪声机制类型（拉普拉斯/高斯）和敏感度等参数。点击"生成推荐"按钮后，系统调用模型服务进行推理，推荐结果以卡片形式展示，包含药物名称、推荐得分、匹配疾病、适应症说明、安全等级、注意事项等详细信息。界面如图5-3所示。',
        'images': [('final_recommendation_1.png', '图5-3  用药推荐页面界面')]
    },
    {
        'name': '患者管理页面',
        'desc': '患者管理页面用于管理患者的基本信息和健康档案数据。页面顶部提供搜索框和筛选功能，用户可按姓名、性别等条件快速查找患者。患者列表以表格形式展示，包含姓名、性别、年龄、联系电话等基本信息。支持新增患者、编辑患者信息和删除患者等操作。新增患者时，需填写姓名、性别、出生日期、联系电话等基本信息，以及慢性疾病、过敏信息、当前用药等健康档案数据。健康档案数据以结构化形式存储，便于推荐系统读取和使用。该页面仅对管理员和医生角色开放。界面如图5-4所示。',
        'images': [
            ('final_patients_1.png', '图5-4a  患者管理页面（列表上部）'),
            ('final_patients_2.png', '图5-4b  患者管理页面（列表下部）'),
        ]
    },
    {
        'name': '隐私配置页面',
        'desc': '隐私配置页面是系统隐私保护功能的核心配置界面，允许用户根据实际需求调整差分隐私参数。页面分为三个主要区域：隐私参数配置区、隐私预算管理区和隐私事件日志区。隐私参数配置区提供隐私预算ε（范围0.01~10.0）、松弛参数δ（默认1e-5）、噪声机制选择（拉普拉斯机制、高斯机制、指数机制）和敏感度设置（默认1.0）等参数的配置。隐私-效用权衡图表直观展示不同ε值下推荐效用（准确率）的变化趋势，帮助用户理解隐私保护强度与推荐效果之间的平衡关系。隐私预算管理区显示当前用户的隐私预算使用情况，包括总预算、已消耗预算和剩余预算。隐私事件日志区记录每次推荐操作的隐私消耗详情，包括时间、消耗量、ε值和推荐类型。界面如图5-5所示。',
        'images': [
            ('final_privacy_1.png', '图5-5a  隐私配置页面（参数配置与效用权衡）'),
            ('final_privacy_2.png', '图5-5b  隐私配置页面（预算管理与事件日志）'),
        ]
    },
    {
        'name': '隐私效果可视化页面',
        'desc': '隐私效果可视化页面提供差分隐私保护效果的直观展示，帮助用户理解隐私机制对推荐结果的影响。页面包含三个视图：总览视图展示隐私预算消耗趋势图和噪声分布对比图，直观呈现隐私预算随推荐次数的消耗变化以及拉普拉斯噪声与高斯噪声的概率密度分布差异；对比分析视图展示不同噪声机制下推荐结果的差异对比，体现差分隐私对推荐排序的影响；深度分析视图展示效用-隐私权衡曲线和特征重要性分析，帮助用户理解不同ε值下推荐准确率的变化趋势。用户可通过交互控件切换视图，实时观察隐私保护效果的变化。界面如图5-6所示。',
        'images': [
            ('final_viz_overview_1.png', '图5-6a  隐私效果可视化页面（总览视图）'),
            ('final_viz_comparison_1.png', '图5-6b  隐私效果可视化页面（对比分析视图）'),
            ('final_viz_analysis_1.png', '图5-6c  隐私效果可视化页面（深度分析视图）'),
        ]
    },
    {
        'name': '管理员仪表盘页面',
        'desc': '管理员仪表盘页面是系统后台管理功能的核心界面，仅对管理员角色开放。页面左侧为管理功能导航栏，包含用户管理、模型训练、系统配置、审计日志等管理模块。用户管理模块展示系统用户列表，支持新增用户、编辑用户角色、删除用户等操作。模型训练模块提供模型训练触发接口和训练状态监控功能。系统配置模块管理全局参数设置。审计日志模块记录系统所有重要操作，包括用户登录、推荐操作、隐私配置变更等，确保系统操作的可追溯性。页面右侧为数据统计概览区域，展示用户总数、推荐次数、隐私事件数等关键指标。界面如图5-7所示。',
        'images': [
            ('final_admin_1.png', '图5-7a  管理员仪表盘（统计概览）'),
            ('final_admin_2.png', '图5-7b  管理员仪表盘（用户管理）'),
        ]
    },
    {
        'name': '权限禁止页面',
        'desc': '权限禁止页面在用户尝试访问超出其角色权限的功能时显示。当非管理员用户尝试访问后台管理等功能时，系统自动重定向至该页面，显示403禁止访问提示信息，明确告知用户当前角色不具备该功能的访问权限。该机制保障了系统的安全性和数据隔离性，确保不同角色的用户只能访问其权限范围内的功能。界面如图5-8所示。',
        'images': [('final_forbidden_1.png', '图5-8  权限禁止页面界面')]
    },
]

doc = Document(doc_path)
body = doc.element.body
target_para = doc.paragraphs[207]
target_element = target_para._element

new_elements = []

for page_info in pages:
    # Description paragraph
    desc_para = doc.add_paragraph(page_info['desc'])
    desc_para.style = doc.styles['Normal']
    pf = desc_para.paragraph_format
    pf.first_line_indent = Cm(0.74)
    pf.line_spacing = 1.5
    pf.space_after = Pt(6)
    for run in desc_para.runs:
        run.font.name = '宋体'
        run.font.size = Pt(12)
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    new_elements.append(desc_para._element)

    # Images with captions
    for img_file, caption in page_info['images']:
        # Image paragraph (centered, 18.25cm width)
        img_path = os.path.join(screenshots_dir, img_file)
        img_para = doc.add_paragraph()
        img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pf2 = img_para.paragraph_format
        pf2.space_before = Pt(4)
        pf2.space_after = Pt(2)
        run = img_para.add_run()
        run.add_picture(img_path, width=Cm(18.25))
        new_elements.append(img_para._element)

        # Caption paragraph
        cap_para = doc.add_paragraph()
        cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pf3 = cap_para.paragraph_format
        pf3.space_after = Pt(6)
        cap_run = cap_para.add_run(caption)
        cap_run.font.name = '宋体'
        cap_run.font.size = Pt(10.5)
        cap_run.font.bold = True
        cap_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        new_elements.append(cap_para._element)

# Move elements from end to after paragraph [207]
for elem in new_elements:
    body.remove(elem)

insert_pos = list(body).index(target_element) + 1
for elem in new_elements:
    body.insert(insert_pos, elem)
    insert_pos += 1

doc.save(output_path)
print(f'Saved to {output_path}')

# Verify
doc2 = Document(output_path)
img_count = 0
for para in doc2.paragraphs:
    drawings = para._element.findall('.//' + qn('w:drawing'))
    if drawings:
        img_count += 1
print(f'Images verified: {img_count}')